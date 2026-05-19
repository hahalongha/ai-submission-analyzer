import csv
import io
import os

import gradio as gr
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd

import config
from ai_engine import analyze_submission, APIKeyError, APIConfigError
from report_generator import generate_markdown_report, export_txt, export_word, export_pdf
from storage import add_record, get_history, get_statistics, delete_record, clear_history, get_categories

matplotlib.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "sans-serif"]
matplotlib.rcParams["axes.unicode_minus"] = False

MAX_CHARS = 5000

PROVIDER_LABELS = {
    "deepseek": "DeepSeek（推荐）",
    "qwen": "通义千问",
    "glm": "智谱GLM（免费额度）",
    "kimi": "Kimi",
}

PROVIDER_CHOICES = [
    f"{k} - {PROVIDER_LABELS.get(k, k)}"
    for k in config.API_CONFIGS.keys()
]


def _parse_provider(label: str) -> str:
    if " - " in label:
        return label.split(" - ")[0]
    return label


def process_submission(text: str, export_format: str, provider_label: str):
    try:
        return _process_submission_impl(text, export_format, provider_label)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"❌ 分析出错：{e}", None, None, None, gr.update(visible=False)


def _process_submission_impl(text: str, export_format: str, provider_label: str):
    if not text.strip():
        return "⚠️ 请输入事件描述内容", None, None, None, gr.update(visible=False)

    if len(text) > MAX_CHARS:
        return f"⚠️ 输入内容超过{MAX_CHARS}字限制，请精简后重试", None, None, None, gr.update(visible=False)

    provider = _parse_provider(provider_label)

    try:
        config.set_provider(provider)
    except ValueError as e:
        msg = f"❌ {e}"
        return msg, None, None, None, gr.update(visible=False), msg, None, None, None, gr.update(visible=False)

    try:
        result = analyze_submission(text.strip())
    except APIKeyError as e:
        return f"🔑 **API密钥错误**\n\n{e}", None, None, None, gr.update(visible=False)
    except APIConfigError as e:
        return f"⚙️ **API配置错误**\n\n{e}", None, None, None, gr.update(visible=False)
    except Exception as e:
        return f"❌ 分析失败：{e}", None, None, None, gr.update(visible=False)

    add_record(result)
    md_report = generate_markdown_report(result)

    txt_path = None
    word_path = None
    pdf_path = None
    has_file = False

    if export_format in ["TXT", "两者都导出", "三者都导出"]:
        try:
            txt_path = export_txt(result)
            has_file = True
        except Exception:
            txt_path = None

    if export_format in ["Word", "两者都导出", "三者都导出"]:
        try:
            word_path = export_word(result)
            has_file = True
        except Exception:
            word_path = None

    if export_format in ["PDF", "三者都导出"]:
        try:
            pdf_path = export_pdf(result)
            has_file = True
        except Exception:
            pdf_path = None

    return md_report, txt_path, word_path, pdf_path, gr.update(visible=has_file)


def update_char_count(text: str):
    count = len(text) if text else 0
    if count > MAX_CHARS:
        return f'<span style="color:#ef4444;font-weight:600;">{count}/{MAX_CHARS} 字（已超限）</span>'
    return f'<span style="color:#94a3b8;">{count}/{MAX_CHARS} 字</span>'


def load_example():
    return "我们学校最近要求所有研究生必须签署一份协议，承诺毕业后三年内不得更换工作单位，否则要赔偿培养费5万元。很多同学觉得这不合理，但辅导员说不签就不给毕业证。这件事在我们学院引起了很大争议，有人已经向教育局投诉了。"


def handle_add_category(name: str, desc: str):
    msg = config.add_category(name, desc)
    cats = list(config.CATEGORY_DESCRIPTIONS.keys())
    first_cat = cats[0] if cats else None
    return (
        msg,
        config.get_categories_text(),
        gr.update(choices=cats, value=first_cat),
        gr.update(value=""),
        gr.update(value=""),
    )


def handle_remove_category(name: str):
    msg = config.remove_category(name)
    cats = list(config.CATEGORY_DESCRIPTIONS.keys())
    first_cat = cats[0] if cats else None
    return (
        msg,
        config.get_categories_text(),
        gr.update(choices=cats, value=first_cat),
        gr.update(value=""),
        gr.update(value=""),
    )


def handle_reset_categories():
    msg = config.reset_categories()
    cats = list(config.CATEGORY_DESCRIPTIONS.keys())
    return (
        msg,
        config.get_categories_text(),
        gr.update(choices=cats, value=cats[0] if cats else None),
        gr.update(value=""),
        gr.update(value=""),
    )


def _build_batch_dataframe(all_results):
    if not all_results:
        return pd.DataFrame(columns=["标题", "分类", "热度", "置信度"]), []
    data = []
    for r in all_results:
        cls_info = r.get("classification", {})
        impact = r.get("impact", {})
        heat = impact.get("heat_score", "?")
        if isinstance(heat, (int, float)):
            if heat >= 8:
                heat_str = f"🔴 {heat}"
            elif heat >= 5:
                heat_str = f"🟡 {heat}"
            else:
                heat_str = f"🟢 {heat}"
        else:
            heat_str = str(heat)
        confidence = float(cls_info.get("confidence", 0))
        data.append([
            r.get("title", "-")[:30],
            cls_info.get("category", "-"),
            heat_str,
            f"{confidence:.0%}",
        ])
    return pd.DataFrame(data, columns=["标题", "分类", "热度", "置信度"]), all_results


def handle_batch_select(evt: gr.SelectData, batch_results_state):
    row_idx = evt.index[0] if isinstance(evt.index, (list, tuple)) else evt.index
    if batch_results_state and 0 <= row_idx < len(batch_results_state):
        result = batch_results_state[row_idx]
        md = generate_markdown_report(result)
        return md
    return ""


def handle_batch_analysis(file_obj, export_format: str, provider_label: str):
    try:
        return _handle_batch_analysis_impl(file_obj, export_format, provider_label)
    except Exception as e:
        import traceback
        traceback.print_exc()
        err_msg = f"❌ 批量分析出错：{e}"
        return (
            gr.update(visible=False),
            [],
            gr.update(value=err_msg, visible=True), None, None, None, gr.update(visible=False),
            "", None, None, None, gr.update(visible=False),
        )


def _handle_batch_analysis_impl(file_obj, export_format: str, provider_label: str):
    _hide_df = gr.update(visible=False)

    if file_obj is None:
        msg = "⚠️ 请先上传投稿文件"
        return _hide_df, [], gr.update(value=msg, visible=True), None, None, None, gr.update(visible=False), "", None, None, None, gr.update(visible=False)

    provider = _parse_provider(provider_label)

    try:
        config.set_provider(provider)
    except ValueError as e:
        msg = f"❌ {e}"
        return _hide_df, [], gr.update(value=msg, visible=True), None, None, None, gr.update(visible=False), "", None, None, None, gr.update(visible=False)

    try:
        content = ""
        if hasattr(file_obj, 'read'):
            raw = file_obj.read()
            if isinstance(raw, bytes):
                content = raw.decode("utf-8", errors="ignore")
            else:
                content = raw
        else:
            with open(file_obj.name if hasattr(file_obj, 'name') else file_obj, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
    except Exception as e:
        msg = f"❌ 读取文件失败：{e}"
        return _hide_df, [], gr.update(value=msg, visible=True), None, None, None, gr.update(visible=False), "", None, None, None, gr.update(visible=False)

    filename = file_obj.name if hasattr(file_obj, 'name') else str(file_obj)
    submissions = []

    if filename.lower().endswith(".csv"):
        try:
            reader = csv.reader(io.StringIO(content))
            for row in reader:
                if row and row[0].strip():
                    submissions.append(row[0].strip())
        except Exception:
            submissions = [line.strip() for line in content.split("\n") if line.strip()]
    else:
        submissions = [line.strip() for line in content.split("\n") if line.strip()]

    if not submissions:
        msg = "⚠️ 文件中没有有效的投稿内容"
        return _hide_df, [], gr.update(value=msg, visible=True), None, None, None, gr.update(visible=False), "", None, None, None, gr.update(visible=False)

    all_results = []
    result_parts = []
    success_count = 0
    fail_count = 0

    for i, sub in enumerate(submissions):
        if len(sub) > MAX_CHARS:
            sub = sub[:MAX_CHARS]
        try:
            result = analyze_submission(sub)
            add_record(result)
            all_results.append(result)
            md = generate_markdown_report(result)
            result_parts.append(md)
            success_count += 1
        except APIKeyError as e:
            result_parts.append(f"## 第{i+1}条 - API密钥错误\n\n{e}")
            fail_count += 1
            break
        except APIConfigError as e:
            result_parts.append(f"## 第{i+1}条 - API配置错误\n\n{e}")
            fail_count += 1
            break
        except Exception as e:
            result_parts.append(f"## 第{i+1}条 - 分析失败\n\n错误：{e}")
            fail_count += 1

    summary_md = (
        f"## 批量分析完成\n\n"
        f"共处理 **{len(submissions)}** 条投稿，成功 **{success_count}** 条，失败 **{fail_count}** 条\n\n"
        f"---\n\n"
    ) + "\n\n---\n\n".join(result_parts)

    batch_df, batch_results_list = _build_batch_dataframe(all_results)

    summary_short = (
        f"✅ 批量分析完成：共 **{len(submissions)}** 条，"
        f"成功 **{success_count}** 条，失败 **{fail_count}** 条\n\n"
        f"👆 点击上方表格中的行查看详细分析结果"
    )

    txt_path = None
    word_path = None
    pdf_path = None
    has_file = False

    if export_format in ["TXT", "两者都导出", "三者都导出"]:
        try:
            os.makedirs("output", exist_ok=True)
            timestamp = __import__('datetime').datetime.now().strftime('%Y%m%d_%H%M%S')
            txt_path = os.path.join("output", f"批量分析结果_{timestamp}.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(summary_md)
            has_file = True
        except Exception:
            txt_path = None

    if export_format in ["Word", "两者都导出", "三者都导出"] and all_results:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            os.makedirs("output", exist_ok=True)
            timestamp = __import__('datetime').datetime.now().strftime('%Y%m%d_%H%M%S')
            word_path = os.path.join("output", f"批量分析报告_{timestamp}.docx")

            doc = Document()
            title_para = doc.add_heading("批量投稿分析报告", level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"共处理 {len(submissions)} 条投稿，成功 {success_count} 条，失败 {fail_count} 条")
            doc.add_paragraph("")

            for i, result in enumerate(all_results):
                doc.add_heading(f"第{i+1}条：{result['title']}", level=1)
                cls_info = result["classification"]
                impact = result["impact"]
                confidence = float(cls_info.get("confidence", 0))
                heat = int(impact.get("heat_score", 5))

                doc.add_heading("事件分类", level=2)
                table = doc.add_table(rows=3, cols=2, style="Light Grid Accent 1")
                table.cell(0, 0).text = "类别"
                table.cell(0, 1).text = cls_info.get("category", "未知")
                table.cell(1, 0).text = "置信度"
                table.cell(1, 1).text = f"{confidence:.0%}"
                table.cell(2, 0).text = "分类理由"
                table.cell(2, 1).text = cls_info.get("reason", "无")

                doc.add_heading("核心摘要", level=2)
                doc.add_paragraph(result["summary"])

                doc.add_heading("影响评估", level=2)
                heat_para = doc.add_paragraph()
                run = heat_para.add_run(f"热度评分：{heat}/10")
                run.bold = True
                run.font.size = Pt(14)
                if heat >= 8:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                elif heat >= 5:
                    run.font.color.rgb = RGBColor(255, 165, 0)
                else:
                    run.font.color.rgb = RGBColor(0, 128, 0)

                impact_table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
                impact_table.cell(0, 0).text = "影响范围"
                impact_table.cell(0, 1).text = impact.get("affected_scope", "未评估")
                impact_table.cell(1, 0).text = "传播潜力"
                impact_table.cell(1, 1).text = impact.get("spread_potential", "未评估")
                impact_table.cell(2, 0).text = "历史相似"
                impact_table.cell(2, 1).text = impact.get("historical_similarity", "未评估")
                impact_table.cell(3, 0).text = "情感强度"
                impact_table.cell(3, 1).text = impact.get("emotional_intensity", "未评估")
                impact_table.cell(4, 0).text = "综合评估"
                impact_table.cell(4, 1).text = impact.get("overall_reason", "未评估")

                doc.add_paragraph("")

            doc.save(word_path)
            has_file = True
        except Exception:
            word_path = None

    if export_format in ["PDF", "三者都导出"] and all_results:
        try:
            from fpdf import FPDF

            os.makedirs("output", exist_ok=True)
            timestamp = __import__('datetime').datetime.now().strftime('%Y%m%d_%H%M%S')
            pdf_path = os.path.join("output", f"批量分析报告_{timestamp}.pdf")

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)

            font_path = ""
            for candidate in ["C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf"]:
                if os.path.exists(candidate):
                    font_path = candidate
                    break

            if font_path:
                pdf.add_font("chinese", "", font_path, uni=True)
                pdf.add_font("chinese", "B", font_path, uni=True)
                font_name = "chinese"
            else:
                font_name = "Helvetica"

            pdf.add_page()
            pdf.set_font(font_name, "B", 20)
            pdf.cell(0, 15, "批量投稿分析报告", ln=True, align="C")
            pdf.ln(3)
            pdf.set_font(font_name, "", 11)
            pdf.cell(0, 7, f"共处理 {len(submissions)} 条投稿，成功 {success_count} 条，失败 {fail_count} 条", ln=True)
            pdf.ln(5)

            for i, result in enumerate(all_results):
                pdf.add_page()
                cls_info = result["classification"]
                impact = result["impact"]
                heat = int(impact.get("heat_score", 5))
                confidence = float(cls_info.get("confidence", 0))

                pdf.set_font(font_name, "B", 16)
                pdf.cell(0, 10, f"第{i+1}条：{result['title']}", ln=True)
                pdf.ln(3)

                pdf.set_font(font_name, "B", 13)
                pdf.cell(0, 8, "事件分类", ln=True)
                pdf.set_font(font_name, "", 11)
                pdf.cell(0, 7, f"类别：{cls_info.get('category', '未知')}", ln=True)
                pdf.cell(0, 7, f"置信度：{confidence:.0%}", ln=True)
                pdf.cell(0, 7, f"分类理由：{cls_info.get('reason', '无')}", ln=True)
                pdf.ln(3)

                pdf.set_font(font_name, "B", 13)
                pdf.cell(0, 8, "核心摘要", ln=True)
                pdf.set_font(font_name, "", 11)
                pdf.multi_cell(0, 7, result["summary"])
                pdf.ln(3)

                pdf.set_font(font_name, "B", 13)
                pdf.cell(0, 8, "影响评估", ln=True)
                pdf.set_font(font_name, "B", 12)
                heat_text = f"热度评分：{heat}/10"
                pdf.cell(0, 7, heat_text, ln=True)
                pdf.set_font(font_name, "", 11)
                pdf.cell(0, 7, f"影响范围：{impact.get('affected_scope', '未评估')}", ln=True)
                pdf.cell(0, 7, f"传播潜力：{impact.get('spread_potential', '未评估')}", ln=True)
                pdf.cell(0, 7, f"历史相似：{impact.get('historical_similarity', '未评估')}", ln=True)
                pdf.cell(0, 7, f"情感强度：{impact.get('emotional_intensity', '未评估')}", ln=True)
                pdf.ln(2)
                pdf.set_font(font_name, "B", 11)
                pdf.cell(0, 7, "综合评估：", ln=True)
                pdf.set_font(font_name, "", 11)
                pdf.multi_cell(0, 7, impact.get("overall_reason", "未评估"))

            pdf.output(pdf_path)
            has_file = True
        except Exception:
            pdf_path = None

    right_hint = "👆 请点击左列表格中的行，查看对应投稿的详细分析结果"

    return (
        gr.update(value=batch_df, visible=True), batch_results_list,
        gr.update(value=summary_short, visible=True), txt_path, word_path, pdf_path, gr.update(visible=has_file),
        right_hint, None, None, None, gr.update(visible=False),
    )


def build_statistics_view():
    stats = get_statistics()

    metrics_md = (
        f'<div class="stats-cards">'
        f'<div class="stat-card"><div class="stat-value">{stats["total"]}</div><div class="stat-label">总投稿数</div></div>'
        f'<div class="stat-card"><div class="stat-value">{stats["today_count"]}</div><div class="stat-label">今日投稿</div></div>'
        f'<div class="stat-card"><div class="stat-value">{stats["high_heat_count"]}</div><div class="stat-label">高热度</div></div>'
        f'<div class="stat-card"><div class="stat-value">{stats["category_count"]}</div><div class="stat-label">分类数</div></div>'
        f'</div>'
    )

    os.makedirs("data", exist_ok=True)
    pie_path = os.path.join("data", "pie_chart.png")
    if stats["category_distribution"]:
        fig, ax = plt.subplots(figsize=(5, 4.5))
        fig.patch.set_alpha(0)
        ax.set_facecolor('none')
        labels = list(stats["category_distribution"].keys())
        sizes = list(stats["category_distribution"].values())
        purple_palette = ['#6b8f71', '#2d4a3e', '#8aab8f', '#3d6b4e', '#a8c4a8',
                          '#f59e0b', '#fbbf24', '#ef4444', '#f87171', '#fca5a5',
                          '#78716c', '#a8a29e', '#d6d3d1']
        colors = [purple_palette[i % len(purple_palette)] for i in range(len(labels))]
        wedges, texts, autotexts = ax.pie(
            sizes, labels=labels, colors=colors, autopct="%1.0f%%",
            startangle=90, textprops={"fontsize": 11},
            pctdistance=0.75, labeldistance=1.12
        )
        for t in autotexts:
            t.set_fontweight("bold")
            t.set_color("#fff")
        ax.set_title("投稿分类分布", fontsize=14, fontweight="bold", pad=12, color='#334155')
        plt.tight_layout()
        fig.savefig(pie_path, dpi=130, bbox_inches="tight", transparent=True)
        plt.close(fig)
    else:
        pie_path = None

    trend_path = os.path.join("data", "trend_chart.png")
    if stats["trend_7days"]:
        fig, ax = plt.subplots(figsize=(7, 4))
        fig.patch.set_alpha(0)
        ax.set_facecolor('none')
        dates = [d["date"][5:] for d in stats["trend_7days"]]
        counts = [d["count"] for d in stats["trend_7days"]]
        gradient_colors = ['#6b8f71', '#7a9d7f', '#8aab8f', '#9aba9f', '#a8c4a8', '#b8d0b8', '#c8dcc8']
        bar_colors = gradient_colors[:len(dates)]
        bars = ax.bar(dates, counts, color=bar_colors, alpha=0.9, width=0.55,
                       edgecolor='#6b8f71', linewidth=0.5)
        for bar, c in zip(bars, counts):
            if c > 0:
                ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.15,
                        str(c), ha="center", va="bottom", fontsize=11, fontweight="bold",
                        color="#6b8f71")
        ax.set_title("近7天投稿趋势", fontsize=14, fontweight="bold", pad=12, color='#334155')
        ax.set_ylabel("投稿数", fontsize=11, color='#64748b')
        ax.set_xlabel("日期", fontsize=11, color='#64748b')
        ax.tick_params(colors='#94a3b8')
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_color('#e2e8f0')
        ax.spines["bottom"].set_color('#e2e8f0')
        ax.set_ylim(bottom=0, top=max(counts) * 1.3 if max(counts) > 0 else 1)
        plt.tight_layout()
        fig.savefig(trend_path, dpi=130, bbox_inches="tight", transparent=True)
        plt.close(fig)
    else:
        trend_path = None

    return metrics_md, pie_path, trend_path


def build_history_dataframe(category_filter: str = "全部", heat_filter: str = "全部"):
    records = get_history()
    if category_filter and category_filter != "全部":
        records = [r for r in records if r.get("category", "未知") == category_filter]
    if heat_filter and heat_filter != "全部":
        if heat_filter == "高热度(8-10)":
            records = [r for r in records if isinstance(r.get("heat_score"), (int, float)) and r["heat_score"] >= 8]
        elif heat_filter == "中热度(5-7)":
            records = [r for r in records if isinstance(r.get("heat_score"), (int, float)) and 5 <= r["heat_score"] < 8]
        elif heat_filter == "低热度(1-4)":
            records = [r for r in records if isinstance(r.get("heat_score"), (int, float)) and r["heat_score"] < 5]

    record_ids = [r.get("id", "") for r in records[:50]]

    if not records:
        return pd.DataFrame(columns=["标题", "分类", "热度", "置信度", "时间"]), record_ids

    data = []
    for r in records[:50]:
        heat = r.get("heat_score", "?")
        if isinstance(heat, (int, float)):
            if heat >= 8:
                heat_str = f"🔴 {heat}"
            elif heat >= 5:
                heat_str = f"🟡 {heat}"
            else:
                heat_str = f"🟢 {heat}"
        else:
            heat_str = str(heat)

        confidence = r.get("confidence", 0)
        data.append([
            r.get("title", "-")[:30],
            r.get("category", "-"),
            heat_str,
            f"{confidence:.0%}",
            r.get("time", "-"),
        ])

    return pd.DataFrame(data, columns=["标题", "分类", "热度", "置信度", "时间"]), record_ids


def build_history_detail_md(record_id: str) -> str:
    if not record_id:
        return ""
    records = get_history()
    for r in records:
        if r.get("id") == record_id:
            heat = r.get("heat_score", "?")
            if isinstance(heat, (int, float)):
                if heat >= 8:
                    heat_label = "🔴 高热度"
                elif heat >= 5:
                    heat_label = "🟡 中热度"
                else:
                    heat_label = "🟢 低热度"
            else:
                heat_label = ""

            confidence = r.get("confidence", 0)

            md = (
                f"### 📋 {r.get('title', '-')}\n\n"
                f"| 项目 | 内容 |\n|------|------|\n"
                f"| 📂 分类 | {r.get('category', '-')} |\n"
                f"| 🔥 热度 | {heat}/10 {heat_label} |\n"
                f"| 📊 置信度 | {confidence:.0%} |\n"
                f"| 🕐 时间 | {r.get('time', '-')} |\n\n"
                f"**🏷️ 分类理由**：{r.get('classification_reason', '无')}\n\n"
                f"**📝 核心摘要**：\n{r.get('summary', '暂无摘要')}\n\n"
                f"**📊 影响评估**：\n"
                f"- 影响范围：{r.get('affected_scope', '未评估')}\n"
                f"- 传播潜力：{r.get('spread_potential', '未评估')}\n"
                f"- 历史相似：{r.get('historical_similarity', '未评估')}\n"
                f"- 情感强度：{r.get('emotional_intensity', '未评估')}\n\n"
                f"**💡 综合评估**：{r.get('overall_reason', '未评估')}"
            )
            return md
    return "⚠️ 未找到该记录"


def handle_save_api_key(provider_label: str, api_key: str):
    provider = _parse_provider(provider_label)
    msg = config.set_api_key(provider, api_key)
    key = config.get_api_key(provider)
    status = config.get_api_key_status(provider)
    status_msg = f"**{config.API_CONFIGS[provider]['name']}** 状态：{status}"
    return msg, config.get_all_key_status(), key, status_msg


def handle_load_key_for_provider(provider_label: str):
    provider = _parse_provider(provider_label)
    key = config.get_api_key(provider)
    status = config.get_api_key_status(provider)
    return key, f"**{config.API_CONFIGS[provider]['name']}** 状态：{status}"


def handle_validate_key(provider_label: str):
    provider = _parse_provider(provider_label)
    return config.validate_api_key(provider)


def handle_refresh_key_status():
    return config.get_all_key_status()


def handle_filter_history(category_filter, heat_filter):
    df, ids = build_history_dataframe(category_filter, heat_filter)
    return df, ids, gr.update(value="", visible=False)


def handle_history_select(evt: gr.SelectData, record_ids):
    row_idx = evt.index[0] if isinstance(evt.index, (list, tuple)) else evt.index
    if record_ids and 0 <= row_idx < len(record_ids):
        selected_id = record_ids[row_idx]
        detail_md = build_history_detail_md(selected_id)
        return selected_id, gr.update(value=detail_md, visible=True)
    return "", gr.update(value="", visible=False)


def handle_delete_history_record(selected_id, category_filter, heat_filter):
    if selected_id:
        delete_record(selected_id.strip())
    new_cats = ["全部"] + get_categories()
    cat_val = category_filter if category_filter in new_cats else "全部"
    df, ids = build_history_dataframe(cat_val, heat_filter)
    return (
        df,
        ids,
        gr.update(choices=new_cats, value=cat_val),
        "",
        gr.update(value="", visible=False),
    )


def handle_clear_all_history():
    clear_history()
    df, ids = build_history_dataframe()
    return (
        df,
        ids,
        gr.update(choices=["全部"], value="全部"),
        gr.update(value="全部"),
        "",
        gr.update(value="", visible=False),
    )


def handle_refresh_history():
    df, ids = build_history_dataframe()
    new_cats = ["全部"] + get_categories()
    return (
        df,
        ids,
        gr.update(choices=new_cats, value="全部"),
        gr.update(value="全部"),
        "",
        gr.update(value="", visible=False),
    )


EXAMPLE_INPUTS = [
    [
        "我们学校最近要求所有研究生必须签署一份协议，承诺毕业后三年内不得更换工作单位，否则要赔偿培养费5万元。很多同学觉得这不合理，但辅导员说不签就不给毕业证。这件事在我们学院引起了很大争议，有人已经向教育局投诉了。"
    ],
    [
        "某互联网公司被曝出实行996工作制，员工每天工作12小时，每周6天，但没有支付任何加班费。有员工因过度劳累住院，公司却以'自愿加班'为由拒绝赔偿。目前已有200多名员工联名向劳动监察部门举报。"
    ],
    [
        "某高校食堂被学生曝光使用过期食材，多名学生出现腹泻症状。学校最初否认，但在学生发布视频证据后，才承认问题并承诺整改。家长和学生对学校处理方式不满，要求第三方调查。"
    ],
]

CUSTOM_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,400;0,500;0,600;0,700;0,800;1,400;1,500&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

@keyframes breatheIn {
    0% { opacity: 0; transform: translateY(18px); }
    100% { opacity: 1; transform: translateY(0); }
}

@keyframes breatheOut {
    0% { opacity: 1; transform: scale(1); }
    100% { opacity: 0; transform: scale(0.96); }
}

@keyframes fadeSlideUp {
    from { opacity: 0; transform: translateY(10px); }
    to { opacity: 1; transform: translateY(0); }
}

@keyframes shimmerLine {
    0% { background-position: -200% 0; }
    100% { background-position: 200% 0; }
}

@keyframes subtleFloat {
    0%, 100% { transform: translateY(0); }
    50% { transform: translateY(-6px); }
}

@keyframes spin {
    from { transform: rotate(0deg); }
    to { transform: rotate(360deg); }
}

:root {
    --parchment: #faf7f2;
    --paper: #ffffff;
    --paper-hover: #fefcf9;
    --ink-primary: oklch(0.18 0.01 90);
    --ink-secondary: oklch(0.38 0.01 90);
    --ink-muted: oklch(0.55 0.01 90);

    --sage: oklch(0.48 0.09 155);
    --sage-light: oklch(0.58 0.08 155);
    --sage-pale: oklch(0.48 0.09 155 / 0.06);
    --sage-soft: oklch(0.48 0.09 155 / 0.12);

    --forest: oklch(0.28 0.04 160);
    --forest-light: oklch(0.38 0.04 160);
    --forest-pale: oklch(0.28 0.04 160 / 0.06);

    --amber: oklch(0.62 0.14 80);
    --amber-light: oklch(0.72 0.12 80);
    --amber-pale: oklch(0.62 0.14 80 / 0.08);

    --mint: oklch(0.52 0.06 165);
    --mint-pale: oklch(0.52 0.06 165 / 0.08);

    --terracotta: oklch(0.55 0.11 40);
    --terracotta-pale: oklch(0.55 0.11 40 / 0.08);

    --warm-shadow: 0 1px 3px rgba(0,0,0,0.03), 0 4px 24px rgba(0,0,0,0.04);
    --warm-shadow-hover: 0 2px 6px rgba(0,0,0,0.04), 0 8px 40px rgba(0,0,0,0.06);
    --elevated-shadow: 0 1px 2px rgba(0,0,0,0.02), 0 8px 48px rgba(0,0,0,0.08);

    --radius-sm: 10px;
    --radius-md: 14px;
    --radius-lg: 18px;

    --ease-out-expo: cubic-bezier(0.16, 1, 0.3, 1);
    --ease-out-quint: cubic-bezier(0.22, 1, 0.36, 1);
    --transition-quick: 180ms var(--ease-out-expo);
    --transition-smooth: 350ms var(--ease-out-expo);
    --transition-slow: 500ms var(--ease-out-expo);

    --panel-height: 680px;
}

* {
    -webkit-font-smoothing: antialiased;
    -moz-osx-font-smoothing: grayscale;
}

body {
    background: var(--parchment) !important;
    min-height: 100vh !important;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    font-size: 16px !important;
    letter-spacing: -0.01em;
    overflow-x: hidden;
}

.gradio-container {
    max-width: 98% !important;
    margin: 0 auto !important;
    padding: 0.6rem 0.9rem !important;
}

.main-header {
    background: linear-gradient(135deg, var(--sage) 0%, var(--forest) 100%);
    padding: 1.4rem 2rem 1.1rem;
    border-radius: var(--radius-lg);
    text-align: center;
    margin-bottom: 0.7rem;
    position: relative;
    overflow: hidden;
    box-shadow: var(--elevated-shadow);
}

.main-header::before {
    content: '';
    position: absolute;
    inset: 0;
    background:
        linear-gradient(90deg, transparent 0%, rgba(255,255,255,0.06) 40%, rgba(255,255,255,0.03) 60%, transparent 100%);
    background-size: 200% 100%;
    animation: shimmerLine 8s ease-in-out infinite;
}

.main-header::after {
    content: '';
    position: absolute;
    top: -50%;
    right: -15%;
    width: 260px;
    height: 260px;
    background: radial-gradient(circle, rgba(255,255,255,0.06) 0%, transparent 70%);
    border-radius: 50%;
}

.gradio-container .main-header h1,
.main-header h1 {
    font-size: 2rem !important;
    font-weight: 700 !important;
    color: #fff !important;
    margin: 0 0 0.2rem !important;
    letter-spacing: -0.015em !important;
    font-family: 'Playfair Display', Georgia, serif !important;
    position: relative;
}

.gradio-container .main-header p,
.main-header p {
    font-size: 1.08rem !important;
    color: rgba(255,255,255,0.9) !important;
    margin: 0 !important;
    font-weight: 400 !important;
    position: relative;
    letter-spacing: 0.02em !important;
}

.header-badges {
    display: flex;
    justify-content: center;
    gap: 0.5rem;
    margin-top: 0.7rem;
    flex-wrap: wrap;
    position: relative;
}

.header-badge {
    padding: 0.3rem 0.9rem;
    border-radius: 9999px;
    font-size: 0.88rem;
    font-weight: 500;
    background: rgba(255,255,255,0.16) !important;
    color: #faf3e0 !important;
    border: 1px solid rgba(255,255,255,0.2) !important;
    transition: all var(--transition-smooth);
    cursor: default;
}

.header-badge:hover {
    background: rgba(255,255,255,0.26) !important;
    transform: translateY(-1px);
}

.main-row {
    display: flex !important;
    gap: 0.7rem !important;
    align-items: stretch !important;
}

.main-row > div {
    align-self: stretch !important;
}

.left-col, .right-col {
    flex: 1 1 0% !important;
    min-width: 0 !important;
    background: var(--paper);
    border: 1px solid oklch(0.60 0.01 90 / 0.15);
    border-radius: var(--radius-lg);
    box-shadow: var(--warm-shadow);
    padding: 0.9rem 1.1rem;
    transition: box-shadow var(--transition-smooth), border-color var(--transition-smooth);
    height: var(--panel-height) !important;
    max-height: var(--panel-height) !important;
    min-height: var(--panel-height) !important;
    overflow: visible !important;
}

.left-col:hover, .right-col:hover {
    box-shadow: var(--warm-shadow-hover);
    border-color: oklch(0.60 0.01 90 / 0.25);
}

.left-col > div, .right-col > div {
    height: 100% !important;
    max-height: 100% !important;
}

.left-col .tabs, .right-col .tabs {
    height: 100% !important;
    display: flex !important;
    flex-direction: column !important;
}

.left-col .tabitem, .right-col .tabitem {
    flex: 1 1 0% !important;
    overflow-y: auto !important;
    overflow-x: visible !important;
    min-height: 0 !important;
    position: relative !important;
}

.left-col .tabitem > div, .right-col .tabitem > div {
    overflow: visible !important;
}

textarea, input[type="text"] {
    background: var(--parchment) !important;
    border: 1px solid oklch(0.60 0.01 90 / 0.18) !important;
    border-radius: var(--radius-sm) !important;
    font-size: 1.02rem !important;
    color: var(--ink-primary) !important;
    line-height: 1.7 !important;
    padding: 0.75rem 0.9rem !important;
    transition: border-color var(--transition-quick), box-shadow var(--transition-quick), background var(--transition-quick);
}

textarea:focus, input[type="text"]:focus {
    border-color: var(--sage-soft) !important;
    box-shadow: 0 0 0 3px var(--sage-pale) !important;
    background: var(--paper) !important;
    outline: none !important;
}

.config-divider {
    margin: 0.65rem 0;
    border: none;
    height: 1px;
    background: linear-gradient(90deg, transparent, oklch(0.60 0.01 90 / 0.2), transparent);
}

.analyze-btn-wrap button {
    width: 100%;
    padding: 0.85rem !important;
    font-size: 1.12rem !important;
    font-weight: 600 !important;
    border-radius: var(--radius-sm) !important;
    background: linear-gradient(135deg, var(--sage) 0%, var(--forest) 100%) !important;
    border: none !important;
    color: #fff !important;
    box-shadow: 0 2px 12px rgba(107, 143, 113, 0.22) !important;
    transition: transform var(--transition-quick), box-shadow var(--transition-smooth);
    letter-spacing: 0.01em;
}

.analyze-btn-wrap button:hover {
    box-shadow: 0 4px 20px rgba(107, 143, 113, 0.32) !important;
    transform: translateY(-2px);
}

.analyze-btn-wrap button:active {
    transform: scale(0.975) translateY(0);
    box-shadow: 0 1px 6px rgba(107, 143, 113, 0.18) !important;
}

button.secondary,
button:not([variant="primary"]):not(.analyze-btn-wrap button):not(.tab-nav button):not(.tab-wrapper button) {
    background: var(--paper) !important;
    border: 1px solid oklch(0.60 0.01 90 / 0.18) !important;
    border-radius: var(--radius-sm) !important;
    color: var(--ink-secondary) !important;
    font-weight: 500 !important;
    font-size: 0.98rem !important;
    transition: all var(--transition-quick);
}

button.secondary:hover,
button:not([variant="primary"]):not(.analyze-btn-wrap button):not(.tab-nav button):not(.tab-wrapper button):hover {
    background: var(--paper-hover) !important;
    border-color: oklch(0.60 0.01 90 / 0.3) !important;
    box-shadow: 0 1px 4px rgba(0,0,0,0.03);
}

button.secondary:active,
button:not([variant="primary"]):not(.analyze-btn-wrap button):not(.tab-nav button):not(.tab-wrapper button):active {
    transform: scale(0.965);
}

.tabs {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    padding: 0 !important;
}

.tab-nav, .tab-wrapper .tab-container {
    background: var(--parchment) !important;
    border-radius: var(--radius-sm) !important;
    padding: 5px !important;
    gap: 4px !important;
    display: flex !important;
    border: 1px solid oklch(0.60 0.01 90 / 0.1);
}

.tab-nav button, .tab-wrapper button {
    border-radius: 7px !important;
    font-size: 1rem !important;
    font-weight: 600 !important;
    color: var(--ink-muted) !important;
    padding: 0.6rem 1.1rem !important;
    transition: all var(--transition-smooth);
    min-height: 44px !important;
    flex: 1 1 0% !important;
    white-space: nowrap !important;
    border: 1px solid transparent !important;
    background: transparent !important;
}

.tab-nav button:hover, .tab-wrapper button:hover {
    background: var(--paper) !important;
    color: var(--ink-secondary) !important;
}

.tab-nav button.selected, .tab-wrapper button.selected {
    background: var(--paper) !important;
    color: var(--sage) !important;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
    border: 1px solid oklch(0.48 0.09 155 / 0.1) !important;
}

.left-tabs .tab-nav button.selected,
.left-tabs .tab-wrapper button.selected {
    color: var(--sage) !important;
    background: var(--paper) !important;
    border: 1px solid oklch(0.48 0.09 155 / 0.12) !important;
}

.right-tabs .tab-nav button.selected,
.right-tabs .tab-wrapper button.selected {
    color: var(--sage) !important;
    background: var(--paper) !important;
    border: 1px solid oklch(0.48 0.09 155 / 0.1) !important;
}

.tabitem {
    border: none !important;
    padding: 0.5rem 0 !important;
}

.gr-file {
    background: var(--parchment) !important;
    border: 1.5px dashed oklch(0.60 0.01 90 / 0.25) !important;
    border-radius: var(--radius-sm) !important;
    transition: all var(--transition-smooth);
}

.gr-file:hover {
    border-color: var(--sage-soft) !important;
    background: var(--paper) !important;
    box-shadow: 0 2px 8px var(--sage-pale);
}

.prose {
    font-size: 1.02rem !important;
    line-height: 1.78 !important;
    color: var(--ink-primary) !important;
}

.prose h1, .prose h2, .prose h3 {
    font-family: 'Playfair Display', Georgia, serif !important;
    color: var(--ink-secondary) !important;
    font-weight: 600 !important;
}

.prose table {
    font-size: 0.98rem !important;
    border-radius: var(--radius-sm) !important;
    overflow: hidden !important;
}

.prose th {
    background: var(--parchment) !important;
    color: var(--ink-secondary) !important;
    font-weight: 600 !important;
    border: 1px solid oklch(0.60 0.01 90 / 0.12) !important;
}

.prose td {
    border: 1px solid oklch(0.60 0.01 90 / 0.08) !important;
}

footer { display: none !important; }

.gr-image {
    border-radius: var(--radius-sm) !important;
    overflow: hidden !important;
}

/* ===== DROPDOWN FIXES ===== */

.history-filter-dropdown {
    position: relative !important;
    z-index: 100 !important;
}

.history-filter-dropdown > div {
    position: relative !important;
}

.history-filter-dropdown [class*="select"] [class*="option"],
.history-filter-dropdown ul[class*="option"],
.history-filter-dropdown [class*="select"] > div:last-child {
    z-index: 10000 !important;
    position: fixed !important;
}

.right-col [class*="dropdown"] [class*="select"] [class*="option"],
.right-col [class*="dropdown"] ul[class*="option"],
.right-col [class*="dropdown"] [class*="select"] > div:last-child {
    z-index: 9999 !important;
    position: fixed !important;
}

label, .svelte-1gfkn6j {
    font-size: 0.95rem !important;
    font-weight: 600 !important;
    color: var(--ink-secondary) !important;
    letter-spacing: 0.01em;
}

input[type="radio"] + span {
    font-size: 0.98rem !important;
}

/* ===== STATISTICS CARDS ===== */

.stats-cards {
    display: grid;
    grid-template-columns: repeat(4, 1fr);
    gap: 0.7rem;
    margin-bottom: 1rem;
}

.stat-card {
    background: var(--parchment);
    border: 1px solid oklch(0.60 0.01 90 / 0.12);
    border-radius: var(--radius-md);
    padding: 0.9rem 0.6rem;
    text-align: center;
    transition: all var(--transition-smooth);
    cursor: default;
}

.stat-card:hover {
    transform: translateY(-3px);
    box-shadow: var(--warm-shadow-hover);
    border-color: var(--sage-soft);
}

.stat-value {
    font-size: 2rem;
    font-weight: 700;
    color: var(--sage);
    line-height: 1.1;
    font-family: 'Playfair Display', Georgia, serif;
    transition: transform var(--transition-quick);
}

.stat-card:hover .stat-value {
    transform: scale(1.06);
}

.stat-label {
    font-size: 0.86rem;
    color: var(--ink-muted);
    font-weight: 500;
    margin-top: 0.3rem;
    text-transform: uppercase;
    letter-spacing: 0.04em;
}

/* ===== HISTORY DATAFRAME ===== */

.history-detail-section {
    margin-top: 0.8rem;
    padding: 1rem;
    background: var(--parchment);
    border: 1px solid oklch(0.60 0.01 90 / 0.12);
    border-radius: var(--radius-md);
}

.history-detail-section h3 {
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 1.2rem;
    font-weight: 700;
    color: var(--ink-primary);
    margin: 0 0 0.6rem;
    line-height: 1.45;
}

button:focus-visible {
    outline: 2px solid var(--sage-soft) !important;
    outline-offset: 2px !important;
}

.loading-spinner {
    display: inline-block;
    width: 18px;
    height: 18px;
    border: 2px solid var(--forest-pale);
    border-top-color: var(--forest);
    border-radius: 50%;
    animation: spin 0.6s linear infinite;
    vertical-align: middle;
    margin-right: 8px;
}

/* ===== SCROLLBAR ===== */

.left-col .tabitem::-webkit-scrollbar,
.right-col .tabitem::-webkit-scrollbar {
    width: 5px;
}

.left-col .tabitem::-webkit-scrollbar-track,
.right-col .tabitem::-webkit-scrollbar-track {
    background: transparent;
    border-radius: 10px;
}

.left-col .tabitem::-webkit-scrollbar-thumb,
.right-col .tabitem::-webkit-scrollbar-thumb {
    background: oklch(0.60 0.01 90 / 0.25);
    border-radius: 10px;
}

.left-col .tabitem::-webkit-scrollbar-thumb:hover,
.right-col .tabitem::-webkit-scrollbar-thumb:hover {
    background: var(--sage-soft);
}

/* ===== RESPONSIVE ===== */

@media (max-width: 768px) {
    .main-row {
        flex-direction: column !important;
    }
    .left-col, .right-col {
        height: auto !important;
        max-height: none !important;
        min-height: 420px !important;
    }
    .stats-cards {
        grid-template-columns: repeat(2, 1fr);
    }
    .main-header h1 {
        font-size: 1.5rem;
    }
    .main-header p {
        font-size: 0.95rem;
    }
}

@media (max-width: 480px) {
    .stats-cards {
        grid-template-columns: 1fr 1fr;
        gap: 0.45rem;
    }
    .stat-value {
        font-size: 1.45rem;
    }
    .history-header {
        flex-wrap: wrap;
    }
}
"""

with gr.Blocks(title="AI辅助投稿分析工具") as demo:

    gr.HTML(
        """
        <style>
            #main-title { color: #fefcf4 !important; }
            #main-subtitle { color: rgba(255,255,255,0.88) !important; }
            .main-badge { color: #fefcf4 !important; }
        </style>
        <div class="main-header">
            <h1 id="main-title" style="color: #fefcf4;">🔍 AI辅助投稿分析工具</h1>
            <p id="main-subtitle" style="color: rgba(255,255,255,0.88);">快速分类 · 智能摘要 · 影响评估 · 一键导出</p>
            <div class="header-badges">
                <span class="header-badge main-badge" style="color: #fefcf4;">🏷️ 自动分类</span>
                <span class="header-badge main-badge" style="color: #fefcf4;">📝 智能摘要</span>
                <span class="header-badge main-badge" style="color: #fefcf4;">📊 热度评估</span>
                <span class="header-badge main-badge" style="color: #fefcf4;">📄 报告导出</span>
                <span class="header-badge main-badge" style="color: #fefcf4;">🔑 多模型Key</span>
            </div>
        </div>
        """
    )

    with gr.Row(equal_height=True, elem_classes="main-row"):
        with gr.Column(scale=5, min_width=420, elem_classes="left-col"):

            with gr.Tabs(elem_classes="left-tabs"):
                with gr.Tab("📝 单条分析"):
                    input_text = gr.Textbox(
                        label="📝 投稿内容",
                        placeholder="请粘贴事件描述文字（支持微信/小红书/抖音内容）...",
                        lines=4,
                        max_lines=12,
                        show_label=True,
                    )

                    char_count_html = gr.HTML(
                        value='<span style="color:#78716c;font-size:0.85rem;">0/5000 字</span>'
                    )

                    with gr.Row():
                        clear_btn = gr.Button("🗑️ 清空", size="sm")
                        example_btn = gr.Button("📌 示例", size="sm")

                    gr.HTML('<hr class="config-divider">')

                    with gr.Row():
                        provider_choice = gr.Radio(
                            choices=PROVIDER_CHOICES,
                            value=PROVIDER_CHOICES[0],
                            label="🤖 AI模型",
                            scale=3,
                        )
                        export_choice = gr.Radio(
                            choices=["仅查看", "TXT", "Word", "两者都导出", "PDF", "三者都导出"],
                            value="两者都导出",
                            label="📄 导出格式",
                            scale=2,
                        )

                    gr.HTML('<div class="analyze-btn-wrap">')
                    analyze_btn = gr.Button("🚀 开始分析", variant="primary", size="lg")
                    gr.HTML("</div>")

                with gr.Tab("📦 批量处理"):
                    batch_file = gr.File(label="上传投稿文件", file_types=[".txt", ".csv"])
                    with gr.Row():
                        batch_provider = gr.Radio(
                            choices=PROVIDER_CHOICES,
                            value=PROVIDER_CHOICES[0],
                            label="🤖 AI模型",
                            scale=3,
                        )
                        batch_export_choice = gr.Radio(
                            choices=["仅查看", "TXT", "Word", "两者都导出", "PDF", "三者都导出"],
                            value="两者都导出",
                            label="📄 导出格式",
                            scale=2,
                        )
                    gr.HTML('<div class="analyze-btn-wrap">')
                    batch_btn = gr.Button("🚀 开始批量分析", variant="primary")
                    gr.HTML("</div>")
                    batch_df = gr.Dataframe(
                        headers=["标题", "分类", "热度", "置信度"],
                        label="批量分析结果（点击行查看详情）",
                        interactive=False,
                        wrap=True,
                        visible=False,
                    )
                    batch_results_state = gr.State([])
                    batch_summary = gr.Markdown(visible=False)
                    with gr.Row(visible=False) as batch_file_row:
                        batch_txt_file = gr.File(label="📄 TXT", scale=1)
                        batch_word_file = gr.File(label="📝 Word", scale=1)
                        batch_pdf_file = gr.File(label="📑 PDF", scale=1)

            gr.HTML("</div>")

        with gr.Column(scale=5, min_width=420, elem_classes="right-col"):

            right_tabs = gr.Tabs(elem_classes="right-tabs")
            with right_tabs:
                with gr.Tab("📋 分析结果"):
                    report_output = gr.Markdown(
                        value="👋 欢迎使用AI辅助投稿分析工具！\n\n请在左侧输入投稿内容，选择AI模型和导出格式后，点击 **🚀 开始分析** 按钮。\n\n分析结果将在此处显示。",
                        show_label=False,
                    )
                    with gr.Row(visible=False) as file_row:
                        txt_file = gr.File(label="📄 TXT", scale=1)
                        word_file = gr.File(label="📝 Word", scale=1)
                        pdf_file = gr.File(label="📑 PDF", scale=1)

                with gr.Tab("📊 数据统计"):
                    stats_btn = gr.Button("🔄 刷新统计数据", size="sm", elem_id="stats_refresh_btn")
                    _init_metrics, _init_pie, _init_trend = build_statistics_view()
                    stats_metrics = gr.HTML(value=_init_metrics, label="核心指标")
                    with gr.Row():
                        stats_pie = gr.Image(value=_init_pie, label="投稿分类分布", scale=1)
                        stats_trend = gr.Image(value=_init_trend, label="近7天投稿趋势", scale=1)

                with gr.Tab("📜 历史记录"):
                    with gr.Row():
                        history_btn = gr.Button("🔄 刷新", size="sm")
                        delete_selected_btn = gr.Button("🗑️ 删除选中", size="sm", variant="stop")
                        clear_all_btn = gr.Button("🗑️ 清空全部", size="sm", variant="stop")
                    with gr.Row():
                        category_filter = gr.Dropdown(
                            choices=["全部"] + get_categories(),
                            value="全部",
                            label="📂 分类筛选",
                            scale=1,
                            allow_custom_value=False,
                            filterable=True,
                            interactive=True,
                            elem_classes="history-filter-dropdown",
                        )
                        heat_filter = gr.Dropdown(
                            choices=["全部", "高热度(8-10)", "中热度(5-7)", "低热度(1-4)"],
                            value="全部",
                            label="🔥 热度筛选",
                            scale=1,
                            allow_custom_value=False,
                            interactive=True,
                            elem_classes="history-filter-dropdown",
                        )
                    _init_df, _init_ids = build_history_dataframe()
                    history_df = gr.Dataframe(
                        value=_init_df,
                        headers=["标题", "分类", "热度", "置信度", "时间"],
                        interactive=True,
                        label="历史记录（点击行查看详情）",
                        wrap=True,
                        static_columns=[0, 1, 2, 3, 4],
                    )
                    record_ids_state = gr.State(value=_init_ids)
                    selected_id_state = gr.State(value="")
                    history_detail = gr.Markdown(value="", label="记录详情", visible=False)

                with gr.Tab("⚙️ 设置"):
                    with gr.Tabs():
                        with gr.Tab("🔑 API密钥"):
                            key_status_display = gr.Markdown(
                                value=config.get_all_key_status(),
                            )
                            refresh_key_btn = gr.Button("🔄 刷新状态", size="sm")
                            gr.HTML('<hr class="config-divider">')
                            key_provider_choice = gr.Radio(
                                choices=PROVIDER_CHOICES,
                                value=PROVIDER_CHOICES[0],
                                label="选择模型",
                            )
                            key_input = gr.Textbox(
                                label="API密钥",
                                placeholder="粘贴对应模型的API Key...",
                                type="password",
                                show_label=True,
                            )
                            key_status_msg = gr.Markdown("")
                            with gr.Row():
                                save_key_btn = gr.Button("💾 保存", variant="primary", size="sm")
                                load_key_btn = gr.Button("📥 加载", size="sm")
                                validate_key_btn = gr.Button("✅ 验证", size="sm")
                            key_action_msg = gr.Markdown("")

                        with gr.Tab("📂 分类管理"):
                            cat_display = gr.Markdown(value=config.get_categories_text())
                            with gr.Row():
                                cat_name = gr.Textbox(label="分类名称", placeholder="如：财经", scale=1)
                                cat_desc = gr.Textbox(label="分类描述", placeholder="涉及金融、股市等", scale=2)
                            with gr.Row():
                                add_cat_btn = gr.Button("➕ 添加", variant="primary", size="sm")
                                reset_cat_btn = gr.Button("🔄 恢复默认", size="sm")
                            cat_remove_name = gr.Dropdown(
                                choices=list(config.CATEGORY_DESCRIPTIONS.keys()),
                                label="选择要删除的分类",
                            )
                            remove_cat_btn = gr.Button("🗑️ 删除分类", variant="stop", size="sm")
                            cat_msg = gr.Markdown("")

    input_text.change(
        fn=update_char_count,
        inputs=input_text,
        outputs=char_count_html,
    )

    clear_btn.click(
        fn=lambda: ("", '<span style="color:#94a3b8;">0/5000 字</span>'),
        outputs=[input_text, char_count_html],
    )

    example_btn.click(
        fn=load_example,
        outputs=input_text,
    )

    analyze_btn.click(
        fn=process_submission,
        inputs=[input_text, export_choice, provider_choice],
        outputs=[report_output, txt_file, word_file, pdf_file, file_row],
    )

    batch_btn.click(
        fn=handle_batch_analysis,
        inputs=[batch_file, batch_export_choice, batch_provider],
        outputs=[batch_df, batch_results_state,
                 batch_summary, batch_txt_file, batch_word_file, batch_pdf_file, batch_file_row,
                 report_output, txt_file, word_file, pdf_file, file_row],
    )

    batch_df.select(
        fn=handle_batch_select,
        inputs=[batch_results_state],
        outputs=[report_output],
    )

    stats_btn.click(
        fn=build_statistics_view,
        outputs=[stats_metrics, stats_pie, stats_trend],
    )

    history_btn.click(
        fn=handle_refresh_history,
        outputs=[history_df, record_ids_state, category_filter, heat_filter, selected_id_state, history_detail],
    )

    history_df.select(
        fn=handle_history_select,
        inputs=[record_ids_state],
        outputs=[selected_id_state, history_detail],
    )

    delete_selected_btn.click(
        fn=handle_delete_history_record,
        inputs=[selected_id_state, category_filter, heat_filter],
        outputs=[history_df, record_ids_state, category_filter, selected_id_state, history_detail],
    )

    category_filter.change(
        fn=handle_filter_history,
        inputs=[category_filter, heat_filter],
        outputs=[history_df, record_ids_state, history_detail],
    )

    heat_filter.change(
        fn=handle_filter_history,
        inputs=[category_filter, heat_filter],
        outputs=[history_df, record_ids_state, history_detail],
    )

    clear_all_btn.click(
        fn=handle_clear_all_history,
        outputs=[history_df, record_ids_state, category_filter, heat_filter, selected_id_state, history_detail],
    )

    add_cat_btn.click(
        fn=handle_add_category,
        inputs=[cat_name, cat_desc],
        outputs=[cat_msg, cat_display, cat_remove_name, cat_name, cat_desc],
    ).then(
        fn=lambda: gr.update(choices=["全部"] + get_categories(), value="全部"),
        outputs=category_filter,
    )

    remove_cat_btn.click(
        fn=handle_remove_category,
        inputs=cat_remove_name,
        outputs=[cat_msg, cat_display, cat_remove_name, cat_name, cat_desc],
    ).then(
        fn=lambda: gr.update(choices=["全部"] + get_categories(), value="全部"),
        outputs=category_filter,
    )

    reset_cat_btn.click(
        fn=handle_reset_categories,
        outputs=[cat_msg, cat_display, cat_remove_name, cat_name, cat_desc],
    ).then(
        fn=lambda: gr.update(choices=["全部"] + get_categories(), value="全部"),
        outputs=category_filter,
    )

    refresh_key_btn.click(
        fn=handle_refresh_key_status,
        outputs=key_status_display,
    )

    key_provider_choice.change(
        fn=handle_load_key_for_provider,
        inputs=key_provider_choice,
        outputs=[key_input, key_status_msg],
    )

    save_key_btn.click(
        fn=handle_save_api_key,
        inputs=[key_provider_choice, key_input],
        outputs=[key_action_msg, key_status_display, key_input, key_status_msg],
    )

    load_key_btn.click(
        fn=handle_load_key_for_provider,
        inputs=key_provider_choice,
        outputs=[key_input, key_status_msg],
    )

    validate_key_btn.click(
        fn=handle_validate_key,
        inputs=key_provider_choice,
        outputs=key_action_msg,
    )


if __name__ == "__main__":
    demo.launch(
        server_name="0.0.0.0",
        server_port=7860,
        js="""
() => {
    function setupTabAutoRefresh() {
        var lastRefresh = 0;
        var COOLDOWN = 3000;
        document.addEventListener('click', function(e) {
            var tabBtn = e.target.closest('.right-tabs [role="tab"]');
            if (!tabBtn) return;
            var now = Date.now();
            if (now - lastRefresh < COOLDOWN) return;
            var text = tabBtn.textContent.trim();
            if (text.includes('数据统计')) {
                lastRefresh = now;
                setTimeout(function() {
                    var btns = document.querySelectorAll('button');
                    for (var i = 0; i < btns.length; i++) {
                        if (btns[i].textContent.trim().includes('刷新统计数据')) {
                            btns[i].click();
                            break;
                        }
                    }
                }, 500);
            } else if (text.includes('历史记录')) {
                lastRefresh = now;
                setTimeout(function() {
                    var btns = document.querySelectorAll('button');
                    for (var i = 0; i < btns.length; i++) {
                        if (btns[i].textContent.trim() === '🔄 刷新') {
                            btns[i].click();
                            break;
                        }
                    }
                }, 500);
            }
        });
    }

    setupTabAutoRefresh();

    function setupClearAllConfirm() {
        document.addEventListener('click', function(e) {
            var btn = e.target.closest('button');
            if (!btn) return;
            var btnText = btn.textContent.trim();
            if (!btnText.includes('清空全部')) return;
            var confirmed = confirm('⚠️ 确定要清空所有历史记录吗？此操作不可撤销！');
            if (!confirmed) {
                e.preventDefault();
                e.stopPropagation();
                e.stopImmediatePropagation();
            }
        }, true);
    }

    setupClearAllConfirm();

    function setupDropdownFix() {
        function repositionDropdownOptions() {
            var dropdowns = document.querySelectorAll('.history-filter-dropdown');
            dropdowns.forEach(function(dd) {
                var optionsList = dd.querySelector('[class*="select"] > div:last-child') ||
                                  dd.querySelector('ul[class*="option"]') ||
                                  dd.querySelector('[class*="option"]') ||
                                  dd.querySelector('div[class*="list"]');
                if (!optionsList) return;
                var isVisible = optionsList.offsetParent !== null ||
                                window.getComputedStyle(optionsList).display !== 'none';
                if (!isVisible) return;
                var rect = dd.getBoundingClientRect();
                optionsList.style.position = 'fixed';
                optionsList.style.left = rect.left + 'px';
                optionsList.style.top = rect.bottom + 'px';
                optionsList.style.width = rect.width + 'px';
                optionsList.style.zIndex = '10000';
            });
        }

        var debounceTimer = null;
        var observer = new MutationObserver(function() {
            clearTimeout(debounceTimer);
            debounceTimer = setTimeout(repositionDropdownOptions, 30);
        });
        observer.observe(document.body, { childList: true, subtree: true, attributes: true });

        document.addEventListener('click', function() {
            setTimeout(repositionDropdownOptions, 50);
            setTimeout(repositionDropdownOptions, 200);
        });
    }

    setupDropdownFix();

    return;
}
""",
        theme=gr.themes.Soft(
            primary_hue=gr.themes.colors.emerald,
            secondary_hue=gr.themes.colors.stone,
            neutral_hue=gr.themes.colors.stone,
            font=gr.themes.GoogleFont("Inter"),
            spacing_size=gr.themes.sizes.spacing_sm,
            radius_size=gr.themes.sizes.radius_lg,
        ),
        css=CUSTOM_CSS,
    )
