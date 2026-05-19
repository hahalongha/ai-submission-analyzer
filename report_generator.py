import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import safe_filename


def generate_markdown_report(result: dict) -> str:
    cls_info = result["classification"]
    impact = result["impact"]
    heat = int(impact.get("heat_score", 5))

    if heat >= 8:
        heat_label = "🔴 高热度"
    elif heat >= 5:
        heat_label = "🟡 中热度"
    else:
        heat_label = "🟢 低热度"

    confidence = float(cls_info.get("confidence", 0))

    md = f"""# 📋 投稿分析报告

---

## 📌 事件标题
{result['title']}

## 🏷️ 事件分类
- **类别**：{cls_info.get('category', '未知')}
- **置信度**：{confidence:.0%}
- **分类理由**：{cls_info.get('reason', '无')}

## 📝 核心摘要
{result['summary']}

## 📊 影响评估

### 热度评分：{heat}/10 {heat_label}

| 评估维度 | 分析结果 |
|---------|---------|
| 影响范围 | {impact.get('affected_scope', '未评估')} |
| 传播潜力 | {impact.get('spread_potential', '未评估')} |
| 历史相似 | {impact.get('historical_similarity', '未评估')} |
| 情感强度 | {impact.get('emotional_intensity', '未评估')} |

### 综合评估
{impact.get('overall_reason', '未评估')}

---
*报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*
"""
    return md


def export_txt(result: dict, output_dir: str = "output") -> str:
    os.makedirs(output_dir, exist_ok=True)
    md_content = generate_markdown_report(result)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = safe_filename(result['title'])
    filename = f"投稿分析_{safe_name}_{timestamp}.txt"
    filepath = os.path.join(output_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md_content)
    return filepath


def export_word(result: dict, output_dir: str = "output") -> str:
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = safe_filename(result['title'])
    filename = f"投稿分析_{safe_name}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()

    title_para = doc.add_heading("投稿分析报告", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("事件标题", level=1)
    doc.add_paragraph(result["title"])

    doc.add_heading("事件分类", level=1)
    cls_info = result["classification"]
    confidence = float(cls_info.get("confidence", 0))
    table = doc.add_table(rows=3, cols=2, style="Light Grid Accent 1")
    table.cell(0, 0).text = "类别"
    table.cell(0, 1).text = cls_info.get("category", "未知")
    table.cell(1, 0).text = "置信度"
    table.cell(1, 1).text = f"{confidence:.0%}"
    table.cell(2, 0).text = "分类理由"
    table.cell(2, 1).text = cls_info.get("reason", "无")

    doc.add_heading("核心摘要", level=1)
    doc.add_paragraph(result["summary"])

    doc.add_heading("影响评估", level=1)
    impact = result["impact"]
    heat = int(impact.get("heat_score", 5))
    heat_para = doc.add_paragraph()
    run = heat_para.add_run(f"热度评分：{heat}/10")
    run.bold = True
    run.font.size = Pt(14)
    if heat >= 8:
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif heat >= 5:
        run.font.color.rgb = RGBColor(255, 165, 0)
    else:
        run.font.color.rgb = RGBColor(0, 128, 0)

    impact_table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    impact_table.cell(0, 0).text = "影响范围"
    impact_table.cell(0, 1).text = impact.get("affected_scope", "未评估")
    impact_table.cell(1, 0).text = "传播潜力"
    impact_table.cell(1, 1).text = impact.get("spread_potential", "未评估")
    impact_table.cell(2, 0).text = "历史相似"
    impact_table.cell(2, 1).text = impact.get("historical_similarity", "未评估")
    impact_table.cell(3, 0).text = "情感强度"
    impact_table.cell(3, 1).text = impact.get("emotional_intensity", "未评估")
    impact_table.cell(4, 0).text = "综合评估"
    impact_table.cell(4, 1).text = impact.get("overall_reason", "未评估")

    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(filepath)
    return filepath


def export_pdf(result: dict, output_dir: str = "output") -> str:
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = safe_filename(result['title'])
    filename = f"投稿分析_{safe_name}_{timestamp}.pdf"
    filepath = os.path.join(output_dir, filename)

    from fpdf import FPDF

    cls_info = result["classification"]
    impact = result["impact"]
    heat = int(impact.get("heat_score", 5))
    confidence = float(cls_info.get("confidence", 0))

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    font_path = _find_chinese_font()
    if font_path:
        pdf.add_font("chinese", "", font_path, uni=True)
        pdf.add_font("chinese", "B", font_path, uni=True)
        font_name = "chinese"
    else:
        font_name = "Helvetica"

    pdf.set_font(font_name, "B", 20)
    pdf.cell(0, 15, "投稿分析报告", ln=True, align="C")
    pdf.ln(5)

    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    pdf.set_font(font_name, "B", 14)
    pdf.cell(0, 10, "事件标题", ln=True)
    pdf.set_font(font_name, "", 12)
    pdf.multi_cell(0, 8, result["title"])
    pdf.ln(3)

    pdf.set_font(font_name, "B", 14)
    pdf.cell(0, 10, "事件分类", ln=True)
    pdf.set_font(font_name, "", 11)
    pdf.cell(0, 7, f"类别：{cls_info.get('category', '未知')}", ln=True)
    pdf.cell(0, 7, f"置信度：{confidence:.0%}", ln=True)
    pdf.cell(0, 7, f"分类理由：{cls_info.get('reason', '无')}", ln=True)
    pdf.ln(3)

    pdf.set_font(font_name, "B", 14)
    pdf.cell(0, 10, "核心摘要", ln=True)
    pdf.set_font(font_name, "", 11)
    pdf.multi_cell(0, 7, result["summary"])
    pdf.ln(3)

    pdf.set_font(font_name, "B", 14)
    pdf.cell(0, 10, "影响评估", ln=True)
    pdf.set_font(font_name, "B", 12)
    if heat >= 8:
        heat_text = f"热度评分：{heat}/10 (高热度)"
    elif heat >= 5:
        heat_text = f"热度评分：{heat}/10 (中热度)"
    else:
        heat_text = f"热度评分：{heat}/10 (低热度)"
    pdf.cell(0, 8, heat_text, ln=True)
    pdf.ln(2)

    pdf.set_font(font_name, "", 11)
    pdf.cell(0, 7, f"影响范围：{impact.get('affected_scope', '未评估')}", ln=True)
    pdf.cell(0, 7, f"传播潜力：{impact.get('spread_potential', '未评估')}", ln=True)
    pdf.cell(0, 7, f"历史相似：{impact.get('historical_similarity', '未评估')}", ln=True)
    pdf.cell(0, 7, f"情感强度：{impact.get('emotional_intensity', '未评估')}", ln=True)
    pdf.ln(2)

    pdf.set_font(font_name, "B", 11)
    pdf.cell(0, 7, "综合评估：", ln=True)
    pdf.set_font(font_name, "", 11)
    pdf.multi_cell(0, 7, impact.get("overall_reason", "未评估"))

    pdf.ln(10)
    pdf.set_font(font_name, "", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 5, f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align="R")

    pdf.output(filepath)
    return filepath


def _find_chinese_font() -> str:
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simfang.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return ""
